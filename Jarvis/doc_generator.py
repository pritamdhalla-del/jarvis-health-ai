from docx import Document

def create_doc(text, filename="output.docx"):

    doc = Document()

    doc.add_heading("JARVIS REPORT", level=1)

    doc.add_paragraph(text)

    doc.save(filename)

    return filename