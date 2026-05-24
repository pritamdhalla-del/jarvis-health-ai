from groq import Groq
from dotenv import load_dotenv
import os
from docx import Document
from tools import *
from doc_generator import create_doc

load_dotenv()

client = Groq(
    api_key=os.getenv("GROQ_API_KEY")
)

def ask_ai(prompt):

    response = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    return response.choices[0].message.content


def run_agent(user_input):

    text = user_input.lower()

    # MULTI ACTION YOUTUBE + DOCUMENT
    if "youtube" in text and "document" in text:

        # Extract song name
        query = text

        query = query.replace("open youtube", "")
        query = query.replace("play", "")
        query = query.replace("and deliver a summary of that song in document format", "")
        query = query.replace("document format", "")
        query = query.strip()

        # Open YouTube
        play_youtube(query)

        # Generate AI Summary
        summary_prompt = f'''
        Give a short emotional summary of the song:
        {query}

        Include:
        - Mood
        - Theme
        - Emotional meaning
        '''

        summary = ask_ai(summary_prompt)

        # Create DOC file
        file = create_doc(summary, "song_summary.docx")

        return f'''
        ✅ YouTube opened for: {query}

        ✅ Summary generated

        ✅ Document created:
        {file}
        '''

    # OPEN CHROME
    elif "open chrome" in text:

        return open_chrome()

    # YOUTUBE ONLY
    elif "youtube" in text:

        query = text.replace("play", "")
        query = query.replace("youtube", "")
        query = query.strip()

        play_youtube(query)

        return f"Playing {query} on YouTube"

    # GOOGLE SEARCH
    elif "google" in text:

        query = text.replace("search google", "")
        query = query.strip()

        search_google(query)

        return f"Searching Google for {query}"

      # GENERAL AI CHAT
    else:

        response = ask_ai(user_input)

        # DOCUMENT GENERATION
        if "document" in text or "doc" in text:

            doc = Document()

            doc.add_heading("JARVIS AI REPORT", 0)

            doc.add_paragraph(response)

            file_path = "jarvis_report.docx"

            doc.save(file_path)

            return f"""
            {response}

            ✅ Document created:
            {file_path}
            """

        return response